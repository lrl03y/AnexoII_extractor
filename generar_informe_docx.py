#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
===============================================================================
📄 GENERADOR DE INFORMES WORD (.docx) CON PYTHON-DOCX
===============================================================================

🎯 DESCRIPCIÓN:
Script que ejecuta consultas SQL en BD Access y genera informes profesionales
en formato Word (.docx) usando la librería python-docx.

🔧 TECNOLOGÍA UTILIZADA:
   ✅ python-docx (EN VEZ DE pywin32/COM)
   ✅ No requiere Microsoft Word instalado
   ✅ Más estable y compatible que win32com

📊 CARACTERÍSTICAS:
   ✅ Conexión opcional a Access vía win32com
   ✅ Generación de tablas formateadas
   ✅ Estilos profesionales
   ✅ Límites de filas para evitar documentos muy largos
   ✅ Información de metadatos (fecha, BD, totales)

💡 REQUISITOS:
   pip install python-docx
   (opcional) pip install pywin32  # solo si quieres conectar a Access

📅 AUTOR: Asistente IA del proyecto AnexoII_extractor
===============================================================================
"""

import os
import sys
from datetime import datetime

# Intentar importar python-docx (REQUERIDO)
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("❌ python-docx no disponible")
    print("💡 Instala con: pip install python-docx")
    sys.exit(1)

# Intentar importar win32com para Access (OPCIONAL)
try:
    import win32com.client
    WIN32_AVAILABLE = True
except ImportError:
    WIN32_AVAILABLE = False
    print("⚠️  win32com no disponible - modo sin conexión a Access")

class DocxReporter:
    """Generador de informes Word con python-docx"""
    
    def __init__(self, db_path="resultados/firmas.accdb"):
        """
        Inicializa el generador de informes
        
        Args:
            db_path (str): Ruta a la base de datos Access
        """
        self.db_path = os.path.abspath(db_path)
        self.access_app = None
        self.db = None
        self.connected = False
        
    def connect_access(self):
        """
        Conecta a la base de datos Access usando win32com
        ⚠️ REQUIERE: Microsoft Access instalado y pywin32
        
        Returns:
            bool: True si conexión exitosa, False en caso contrario
        """
        if not WIN32_AVAILABLE:
            return False
            
        try:
            if not os.path.exists(self.db_path):
                print(f"❌ Base de datos no encontrada: {self.db_path}")
                return False
                
            self.access_app = win32com.client.Dispatch("Access.Application")
            self.access_app.OpenCurrentDatabase(self.db_path)
            self.db = self.access_app.CurrentDb()
            self.connected = True
            print(f"✅ Conectado a Access: {self.db_path}")
            return True
        except Exception as e:
            print(f"❌ Error conectando a Access: {e}")
            return False
    
    def disconnect_access(self):
        """
        Desconecta limpiamente de Access
        """
        try:
            if self.access_app:
                self.access_app.CloseCurrentDatabase()
                self.access_app.Quit()
        except:
            pass
        self.connected = False
    
    def ejecutar_consulta(self, sql):
        """
        Ejecuta una consulta SQL en la base de datos Access
        
        Args:
            sql (str): Consulta SQL a ejecutar
            
        Returns:
            dict: Diccionario con 'campos' y 'resultados', o None si error
        """
        if not self.connected:
            print("❌ No hay conexión a Access")
            return None
            
        try:
            print("🔍 Ejecutando consulta SQL...")
            recordset = self.db.OpenRecordset(sql)
            
            # Obtener nombres de campos
            campos = []
            for i in range(recordset.Fields.Count):
                campos.append(recordset.Fields(i).Name)
            
            # Obtener datos de resultados
            resultados = []
            if not recordset.BOF:  # Si hay datos
                recordset.MoveFirst()
                while not recordset.EOF:
                    fila = []
                    for i in range(recordset.Fields.Count):
                        valor = recordset.Fields(i).Value
                        if valor is None:
                            fila.append("")
                        else:
                            fila.append(str(valor))
                    resultados.append(fila)
                    recordset.MoveNext()
            
            recordset.Close()
            print(f"✅ Consulta ejecutada: {len(resultados)} filas encontradas")
            return {'campos': campos, 'resultados': resultados}
        except Exception as e:
            print(f"❌ Error ejecutando consulta: {e}")
            return None
    
    def crear_informe_docx(self, titulo, resultado, nombre_archivo):
        """
        Crea un informe profesional en formato Word (.docx)
        
        Args:
            titulo (str): Título del informe
            resultado (dict): Diccionario con 'campos' y 'resultados'
            nombre_archivo (str): Nombre del archivo de salida (sin .docx)
            
        Returns:
            bool: True si éxito, False si error
        """
        try:
            print("📝 Creando informe Word con python-docx...")
            
            # Crear documento nuevo
            doc = Document()
            
            # Título principal con formato
            heading = doc.add_heading(titulo, 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Información del informe
            info_para = doc.add_paragraph()
            run_fecha = info_para.add_run(f'Generado el: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
            run_fecha.bold = True
            doc.add_paragraph(f'Base de datos: {os.path.basename(self.db_path)}')
            doc.add_paragraph(f'Total de filas: {len(resultado["resultados"])}')
            
            # Línea separadora
            doc.add_paragraph()
            doc.add_paragraph('=' * 80)
            doc.add_paragraph()
            
            # Tabla de resultados
            if resultado['resultados']:
                # Crear tabla con estilo de rejilla
                tabla = doc.add_table(
                    rows=1, 
                    cols=len(resultado['campos']),
                    style='Table Grid'
                )
                
                # Cabecera de tabla con formato
                hdr_cells = tabla.rows[0].cells
                for i, campo in enumerate(resultado['campos']):
                    hdr_cells[i].text = str(campo)
                    # Formato de texto en cabecera
                    for paragraph in hdr_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                
                # Datos de la tabla (limitar a 100 filas)
                max_filas = min(len(resultado['resultados']), 100)
                for fila in resultado['resultados'][:max_filas]:
                    row_cells = tabla.add_row().cells
                    for i, valor in enumerate(fila):
                        if i < len(row_cells):
                            # Limitar longitud de texto para evitar celdas muy anchas
                            row_cells[i].text = str(valor)[:250]
                
                # Nota informativa si hay más filas
                if len(resultado['resultados']) > 100:
                    doc.add_paragraph()
                    doc.add_paragraph(
                        f'ℹ️  Solo se muestran las primeras 100 filas. '
                        f'Total en base de datos: {len(resultado["resultados"])} filas.'
                    ).italic = True
            
            else:
                # Mensaje cuando no hay resultados
                doc.add_paragraph("📊 No se encontraron resultados para esta consulta.")
            
            # Pie de página con información adicional
            doc.add_page_break()
            doc.add_paragraph("📄 Informe generado automáticamente por AnexoII Extractor")
            doc.add_paragraph(f"📅 Fecha: {datetime.now().strftime('%d/%m/%Y')}")
            
            # Guardar documento
            if not nombre_archivo.endswith('.docx'):
                nombre_archivo += '.docx'
            
            ruta_completa = os.path.join("resultados", nombre_archivo)
            doc.save(ruta_completa)
            
            print(f"✅ Informe generado exitosamente: {ruta_completa}")
            print("💡 Puedes abrirlo directamente con Word, LibreOffice Writer, Google Docs, etc.")
            return True
            
        except Exception as e:
            print(f"❌ Error creando informe Word: {e}")
            return False

def mostrar_menu_principal():
    """Muestra el menú principal con opciones"""
    print("\n" + "📄 MENÚ PRINCIPAL".center(50, "="))
    print("1. 📊 Ejecutar consulta SQL y generar informe")
    print("2. 📋 Crear informe con datos de ejemplo")
    print("3. 💾 Solo guardar consulta SQL en archivo")
    print("4. ❌ Salir")
    print("=" * 50)

def main():
    """Función principal del script"""
    print("📄 GENERADOR DE INFORMES WORD CON PYTHON-DOCX")
    print("🔒 NO REQUIERE Microsoft Word instalado")
    print("=" * 60)
    
    # Crear carpeta de resultados si no existe
    if not os.path.exists("resultados"):
        try:
            os.makedirs("resultados")
            print("📁 Carpeta 'resultados' creada")
        except Exception as e:
            print(f"❌ Error creando carpeta 'resultados': {e}")
            return
    
    # Inicializar reporter
    reporter = DocxReporter()
    
    # Conectar a Access (opcional)
    print("\n🔌 Intentando conectar a Access...")
    if reporter.connect_access():
        print("✅ Modo completo: Conexión a Access disponible")
    else:
        print("⚠️  Modo limitado: Sin conexión a Access")
        print("💡 Puedes crear informes con datos manuales o archivos")
    
    try:
        while True:
            mostrar_menu_principal()
            opcion = input("\n🔢 Selecciona opción (1-4): ").strip()
            
            if opcion == "1" and reporter.connected:
                # Ejecutar consulta en Access y generar informe
                print("\n" + "🔍 EJECUTAR CONSULTA EN ACCESS".center(40, "-"))
                print("📝 Introduce consulta SQL (una línea, sin comentarios):")
                sql = input("🔍 SQL: ").strip()
                
                if sql:
                    # Limpiar comentarios y asegurar punto y coma
                    sql = sql.split('--')[0].strip()
                    if not sql.endswith(';'):
                        sql += ';'
                    
                    # Ejecutar consulta
                    resultado = reporter.ejecutar_consulta(sql)
                    if resultado:
                        titulo = input("📄 Título del informe: ").strip()
                        if not titulo:
                            titulo = "Informe de Datos"
                        
                        nombre_archivo = input("💾 Nombre archivo (sin .docx): ").strip()
                        if not nombre_archivo:
                            nombre_archivo = f"informe_{int(datetime.now().timestamp())}"
                        
                        print("\n" + "🔄 GENERANDO INFORME...".center(40, "-"))
                        if reporter.crear_informe_docx(titulo, resultado, nombre_archivo):
                            print("🎉 ¡Informe generado correctamente!")
                        else:
                            print("❌ Error generando informe")
                    else:
                        print("❌ Error ejecutando consulta en Access")
                        
            elif opcion == "2":
                # Crear informe con datos de ejemplo
                print("\n" + "📋 INFORME DE EJEMPLO".center(40, "-"))
                
                # Datos de ejemplo realistas
                resultado = {
                    'campos': ['Ayuntamiento', 'Total_Gastos', 'Importe_Promedio', 'Max_Importe'],
                    'resultados': [
                        ['Murcia', '25', '15000', '100000'],
                        ['Abanilla', '18', '12500', '85000'],
                        ['Blanca', '15', '11000', '75000'],
                        ['Bullas', '12', '9800', '65000'],
                        ['Aledo', '10', '8500', '55000']
                    ]
                }
                
                titulo = input("📄 Título del informe: ").strip()
                if not titulo:
                    titulo = "Informe de Ejemplo - Resumen de Gastos"
                
                nombre_archivo = input("💾 Nombre archivo (sin .docx): ").strip()
                if not nombre_archivo:
                    nombre_archivo = "informe_ejemplo"
                
                print("\n" + "🔄 GENERANDO INFORME DE EJEMPLO...".center(40, "-"))
                if reporter.crear_informe_docx(titulo, resultado, nombre_archivo):
                    print("🎉 ¡Informe de ejemplo generado correctamente!")
                else:
                    print("❌ Error generando informe de ejemplo")
                    
            elif opcion == "3":
                # Solo guardar consulta SQL en archivo
                print("\n" + "💾 GUARDAR CONSULTA SQL".center(40, "-"))
                print("📝 Introduce consulta SQL para guardar:")
                sql = input("🔍 SQL: ").strip()
                
                if sql:
                    nombre_archivo = input("💾 Nombre archivo SQL (sin .sql): ").strip()
                    if not nombre_archivo:
                        nombre_archivo = f"consulta_{int(datetime.now().timestamp())}"
                    if not nombre_archivo.endswith('.sql'):
                        nombre_archivo += '.sql'
                    
                    ruta_sql = os.path.join("resultados", nombre_archivo)
                    try:
                        with open(ruta_sql, 'w', encoding='utf-8') as f:
                            f.write(sql.strip() + ';\n')
                        print(f"✅ Consulta guardada en: {ruta_sql}")
                        print("💡 Puedes ejecutarla después desde Access manualmente")
                    except Exception as e:
                        print(f"❌ Error guardando consulta: {e}")
                
            elif opcion == "4":
                print("\n" + "👋 ¡GRACIAS POR USAR EL GENERADOR!".center(50, "="))
                print("💡 Recuerda: Los informes .docx se pueden abrir con:")
                print("   • Microsoft Word")
                print("   • LibreOffice Writer") 
                print("   • Google Docs")
                print("   • Cualquier procesador de textos moderno")
                print("=" * 50)
                break
                
            else:
                print("❌ Opción no válida. Por favor, selecciona 1, 2, 3 o 4.")
                
    except KeyboardInterrupt:
        print("\n\n👋 ¡Hasta luego! (Interrumpido por el usuario)")
    except Exception as e:
        print(f"❌ Error general no controlado: {e}")
        print("💡 Por favor, reporta este error con los detalles anteriores")
    finally:
        # Cerrar conexiones limpiamente
        print("\n" + "🧹 CERRANDO CONEXIONES...".center(40, "-"))
        reporter.disconnect_access()

if __name__ == "__main__":
    main()